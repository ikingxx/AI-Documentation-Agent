import os
from docx import Document

def generate_report(summarized_text, output_dir="outputs", report_name="summary_report"):
    try:
        # Make sure the output directory exists
        os.makedirs(output_dir, exist_ok=True)

        # TEXT REPORT
        txt_file_path = os.path.join(output_dir, f"{report_name}.txt")
        with open(txt_file_path, "w", encoding="utf-8") as txt_file:
            txt_file.write(summarized_text)

        # WORD REPORT
        docx_file_path = os.path.join(output_dir, f"{report_name}.docx")
        doc = Document()
        doc.add_heading('Document Summary Report', 0)
        doc.add_paragraph(summarized_text)
        doc.save(docx_file_path)

        print(f"✅ Reports generated successfully!\n- {txt_file_path}\n- {docx_file_path}")

        return txt_file_path, docx_file_path

    except Exception as e:
        print(f"❌ Error generating report: {e}")
        return None, None
